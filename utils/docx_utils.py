from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def parse_docx(path_or_bytes):
    if hasattr(path_or_bytes, "read"):
        doc = Document(path_or_bytes)
    else:
        doc = Document(path_or_bytes)
    texts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            texts.append(p.text.strip())
    return "\\n".join(texts)

def detect_document_type(text):
    txt = text.lower()
    if "articles of association" in txt or ("articles" in txt and "association" in txt):
        return "Articles of Association"
    if "memorandum of association" in txt or "memorandum" in txt:
        return "Memorandum of Association"
    if "ubo" in txt or "ultimate beneficial owner" in txt:
        return "UBO Declaration Form"
    if "register of members" in txt:
        return "Register of Members and Directors"
    if "resolution" in txt:
        return "Board Resolution / Shareholder Resolution"
    return "Unknown Document Type"

def insert_comments_and_save(input_path, output_path, issues):
    doc = Document(input_path)
    doc.add_page_break()
    doc.add_paragraph("=== Automated Review Comments ===")
    for i, iss in enumerate(issues, start=1):
        p = doc.add_paragraph(f"{i}. [{iss.get('severity','Medium')}] {iss.get('issue')} (Suggestion: {iss.get('suggestion','N/A')})")
        # include citations if present
        cites = iss.get("citations", [])
        if cites:
            doc.add_paragraph(f"   Citations (by KB passage index): {cites}")
    doc.save(output_path)